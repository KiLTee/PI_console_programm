import docx
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# from datetime import datetime


def print_metodic(TABLE_INF, key, author_name, path) -> None:
    doc = docx.Document()
    p = doc.add_paragraph()
    p1 = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # date = str(datetime.now().day) + '.' + str(datetime.now().month) + '.' + str(datetime.now().year)
    # time = str(datetime.now().hour) + ':' + str(datetime.now().minute)
    run = p.add_run(f"Паспорт уязвимости {key} ")
    run1 = p1.add_run(f"Автор отчета: {author_name}")

    run.font.name = "TimesNewRoman"
    run.font.size = Pt(14)
    run1.font.name = "TimesNewRoman"
    run1.font.size = Pt(14)
    name = key[4:]

    print_table(doc, TABLE_INF)
    doc.save(f"{path}/Паспорт уязвимости {name}.docx")


def print_table(doc, TABLE_INF) -> None:
    table = doc.add_table(
        rows=(len(TABLE_INF)), cols=2
    )  # Добавляем таблицу размером RxC
    table.style = "Table Grid"  # Добавляем рамку в таблице
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    row = 0
    for key in TABLE_INF:
        cell = table.cell(row, 0)
        cell.text = key
        cell = table.cell(row, 1)
        cell.text = str(TABLE_INF[key])

        change_font_cell(table, row, 0)
        change_font_cell(table, row, 1)

        row += 1

def change_font_cell(table, row, col=0) -> None:
    paragraphs = table.cell(row, col).paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.font.name = "TimesNewRoman"
